import os
import json
import asyncio
from datetime import datetime
from typing import Dict, Any, Optional, Union, List
from dataclasses import dataclass, asdict
import logging
import requests
from docx.shared import RGBColor

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

# Web framework
from flask import Flask, render_template, request, jsonify, session, send_file
import io

# Document generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# AI and search tools
from langchain_openai import AzureChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from tavily import TavilyClient
from pydantic import BaseModel, Field, validator

# LangGraph
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import ToolNode
from langchain_core.tools import tool

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
OPENAI_API_VERSION = os.getenv("OPENAI_API_VERSION")
deployment_name = os.getenv("deployment_name")
model_name = os.getenv("model_name")

# Load search API keys from environment variables
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "YOUR_TAVILY_API_KEY")
SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY", "YOUR_SERPAPI_API_KEY")
BRAVE_SEARCH_API_KEY = os.getenv("BRAVE_SEARCH_API_KEY", "YOUR_BRAVE_SEARCH_API_KEY")

# Helper function to convert various types to string
def to_string(value: Union[str, int, float, List, None]) -> str:
    """Convert various types to string safely"""
    if value is None:
        return "Not found"
    elif isinstance(value, str):
        return value
    elif isinstance(value, (int, float)):
        return str(value)
    elif isinstance(value, list):
        return "; ".join([str(item) for item in value if item is not None])
    else:
        return str(value)

# Enhanced Multi-tool search implementation
class EnhancedMultiToolSearchEngine:
    """Enhanced search engine with guaranteed multi-source usage"""
    
    def __init__(self):
        self.tavily_client = None
        self.serpapi_key = SERPAPI_API_KEY
        self.brave_key = BRAVE_SEARCH_API_KEY
        self.search_sources_used = []
        
        if TAVILY_API_KEY and TAVILY_API_KEY != "YOUR_TAVILY_API_KEY":
            try:
                self.tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                logger.info("Tavily client initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize Tavily client: {str(e)}")
    
    def _search_tavily(self, query: str, max_results: int = 15) -> List[Dict]:
        """Primary search using Tavily API"""
        if not self.tavily_client:
            return []
        
        try:
            response = self.tavily_client.search(
                query=query,
                search_depth="advanced",
                max_results=max_results,
                include_domains=[
                    "bloomberg.com", "reuters.com", "crunchbase.com", 
                    "linkedin.com", "sec.gov", "marketwatch.com", 
                    "forbes.com", "yahoo.com", "cnbc.com", "finance.yahoo.com",
                    "businesswire.com", "prnewswire.com", "companiesmarketcap.com",
                    "glassdoor.com", "indeed.com", "zoominfo.com"
                ]
            )
            
            results = []
            for result in response.get('results', []):
                content = result.get('content', '')
                if content and len(content.strip()) > 20:
                    results.append({
                        'title': result.get('title', ''),
                        'content': content,
                        'url': result.get('url', ''),
                        'score': result.get('score', 0),
                        'source': 'tavily'
                    })
            
            logger.info(f"Tavily returned {len(results)} quality results")
            return results
            
        except Exception as e:
            logger.error(f"Tavily search error: {str(e)}")
            return []
    
    def _search_serpapi(self, query: str, num_results: int = 20) -> List[Dict]:
        """Fallback search using SerpApi"""
        if not self.serpapi_key or self.serpapi_key == "YOUR_SERPAPI_API_KEY":
            return []
        
        try:
            url = "https://serpapi.com/search.json"
            params = {
                'q': query,
                'api_key': self.serpapi_key,
                'num': num_results,
                'gl': 'us',
                'hl': 'en',
                'safe': 'off'
            }
            
            response = requests.get(url, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            
            results = []
            organic_results = data.get('organic_results', [])
            
            for idx, result in enumerate(organic_results):
                snippet = result.get('snippet', '')
                if snippet and len(snippet.strip()) > 20:
                    results.append({
                        'title': result.get('title', ''),
                        'content': snippet,
                        'url': result.get('link', ''),
                        'score': len(organic_results) - idx + 50,
                        'source': 'serpapi'
                    })
            
            if 'knowledge_graph' in data:
                kg = data['knowledge_graph']
                description = kg.get('description', '')
                if description and len(description.strip()) > 20:
                    results.append({
                        'title': f"Knowledge Graph: {kg.get('title', '')}",
                        'content': description,
                        'url': kg.get('website', ''),
                        'score': 100,
                        'source': 'serpapi_kg'
                    })
            
            logger.info(f"SerpApi returned {len(results)} quality results")
            return results
            
        except Exception as e:
            logger.error(f"SerpApi search error: {str(e)}")
            return []
    
    def _search_brave(self, query: str, count: int = 20) -> List[Dict]:
        """Secondary fallback using Brave Search API"""
        if not self.brave_key or self.brave_key == "YOUR_BRAVE_SEARCH_API_KEY":
            return []
        
        try:
            url = "https://api.search.brave.com/res/v1/web/search"
            headers = {
                'Accept': 'application/json',
                'Accept-Encoding': 'gzip',
                'X-Subscription-Token': self.brave_key
            }
            
            # Limit count to max 20 to avoid 422 errors
            actual_count = min(count, 20)
            
            params = {
                'q': query,
                'count': actual_count,
                'search_lang': 'en',
                'country': 'US'
            }
            
            response = requests.get(url, headers=headers, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            
            results = []
            web_results = data.get('web', {}).get('results', [])
            
            for idx, result in enumerate(web_results):
                description = result.get('description', '')
                if description and len(description.strip()) > 20:
                    results.append({
                        'title': result.get('title', ''),
                        'content': description,
                        'url': result.get('url', ''),
                        'score': len(web_results) - idx + 30,
                        'source': 'brave'
                    })
            
            logger.info(f"Brave Search returned {len(results)} quality results")
            return results
            
        except Exception as e:
            logger.error(f"Brave Search error: {str(e)}")
            return []
    
    def comprehensive_multi_search(self, query: str, mode: str = "quick") -> Dict[str, Any]:
        """
        Comprehensive multi-tool search with mode support
        mode: 'quick' for standard search, 'deep' for exhaustive search
        """
        logger.info(f"Starting {mode} search for: {query}")
        
        # Adjust search parameters based on mode
        if mode == "deep":
            tavily_max = 30
            serpapi_num = 40
            brave_count = 40
            top_results = 50
        else:
            tavily_max = 15
            serpapi_num = 20
            brave_count = 20
            top_results = 25
        
        all_results = []
        sources_used = []
        
        engines = [
            ("tavily", lambda: self._search_tavily(query, tavily_max)),
            ("serpapi", lambda: self._search_serpapi(query, serpapi_num)), 
            ("brave", lambda: self._search_brave(query, brave_count))
        ]
        
        for engine_name, search_func in engines:
            try:
                engine_results = search_func()
                if engine_results:
                    all_results.extend(engine_results)
                    sources_used.append(engine_name)
                    logger.info(f"{engine_name}: {len(engine_results)} results")
            except Exception as e:
                logger.error(f"{engine_name} failed: {str(e)}")
        
        # Deduplicate and filter
        seen_urls = set()
        unique_results = []
        
        for result in all_results:
            url = result.get('url', '').lower().strip()
            content = result.get('content', '').strip()
            
            if not content or len(content) < 20:
                continue
                
            simplified_url = url.replace('www.', '').replace('http://', '').replace('https://', '').split('?')[0]
            
            if simplified_url not in seen_urls and simplified_url:
                seen_urls.add(simplified_url)
                unique_results.append(result)
        
        unique_results.sort(key=lambda x: x.get('score', 0), reverse=True)
        unique_results = unique_results[:top_results]
        
        date_freshness = f"Data collected on {datetime.now().strftime('%Y-%m-%d')}"
        
        return {
            'results': unique_results,
            'sources_used': sources_used,
            'date_freshness': date_freshness,
            'total_results': len(unique_results)
        }

# Initialize search engine
multi_search_engine = EnhancedMultiToolSearchEngine()

# Pydantic models for Quick and Deep Search
class QuickCompanyResearch(BaseModel):
    """Quick Search - 6-7 page equivalent report"""
    company: str
    employee_count: str
    annual_revenue: str
    sic_codes: str
    naics_codes: str
    company_official_website: str
    address: str
    phone_number: str
    recent_funding: str
    current_erp: str
    leadership_changes: str
    recent_news: str
    recent_sap_job_postings: str
    strengths: str
    weaknesses: str
    opportunities: str
    threats: str
    founded_year: str = "Not found"
    ceo_name: str = "Not found"
    industry_sector: str = "Not found"
    stock_symbol: str = "Not found"
    market_cap: str = "Not found"
    headquarters_location: str = "Not found"
    business_model: str = "Not found"
    competitors: str = "Not found"
    research_timestamp: str = Field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    search_sources_used: str = "Multiple sources"
    data_freshness: str = "Recent data"
    search_mode: str = "Quick Search"

    @validator('*', pre=True)
    def convert_to_string(cls, v):
        return to_string(v)

    class Config:
        extra = "ignore"

class DeepCompanyResearch(BaseModel):
    """Deep Search - 15-20 page equivalent report with citations"""
    # All fields from Quick Search plus additional deep analysis fields
    company: str
    employee_count: str
    annual_revenue: str
    sic_codes: str
    naics_codes: str
    company_official_website: str
    address: str
    phone_number: str
    recent_funding: str
    current_erp: str
    leadership_changes: str
    board_members: str
    recent_news: str
    recent_sap_job_postings: str
    strengths: str
    weaknesses: str
    opportunities: str
    threats: str
    founded_year: str = "Not found"
    ceo_name: str = "Not found"
    industry_sector: str = "Not found"
    stock_symbol: str = "Not found"
    market_cap: str = "Not found"
    headquarters_location: str = "Not found"
    business_model: str = "Not found"
    competitors: str = "Not found"
    
    # Deep search exclusive fields
    detailed_financials: str = "Not found"
    product_portfolio: str = "Not found"
    customer_segments: str = "Not found"
    market_position_analysis: str = "Not found"
    technology_stack: str = "Not found"
    partnerships_alliances: str = "Not found"
    regulatory_compliance: str = "Not found"
    sustainability_initiatives: str = "Not found"
    innovation_rd: str = "Not found"
    geographic_presence: str = "Not found"
    competitive_advantages: str = "Not found"
    risk_factors: str = "Not found"
    growth_strategy: str = "Not found"
    digital_transformation: str = "Not found"
    talent_acquisition: str = "Not found"
    
    # Citations
    sources_citations: str = "Not found"
    
    research_timestamp: str = Field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    search_sources_used: str = "Multiple sources"
    data_freshness: str = "Recent data"
    search_mode: str = "Deep Search"

    @validator('*', pre=True)
    def convert_to_string(cls, v):
        return to_string(v)

    class Config:
        extra = "ignore"

# Agent State
class AgentState(BaseModel):
    company_name: str
    search_mode: str = "quick"  # 'quick' or 'deep'
    search_results: Dict[str, Any] = Field(default_factory=dict)
    research_data: Optional[Union[QuickCompanyResearch, DeepCompanyResearch]] = None
    current_step: str = "start"
    error: Optional[str] = None
    search_metadata: Dict[str, Any] = Field(default_factory=dict)

    class Config:
        arbitrary_types_allowed = True

# Initialize Azure OpenAI
llm = AzureChatOpenAI(
    api_key=OPENAI_API_KEY,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_version=OPENAI_API_VERSION,
    deployment_name=deployment_name,
    model_name=model_name,
    temperature=0.1
)

# Tools
@tool
def search_company_info(query: str, mode: str = "quick") -> str:
    """Search for company information"""
    try:
        search_result = multi_search_engine.comprehensive_multi_search(query, mode)
        
        formatted_results = []
        for result in search_result['results']:
            formatted_results.append({
                'title': result.get('title', ''),
                'content': result.get('content', ''),
                'url': result.get('url', ''),
                'source_engine': result.get('source', ''),
                'score': result.get('score', 0)
            })
        
        metadata = {
            'search_engines_used': search_result['sources_used'],
            'date_freshness': search_result['date_freshness'],
            'total_results_found': search_result['total_results']
        }
        
        return json.dumps({
            'results': formatted_results,
            'metadata': metadata
        }, indent=2)
        
    except Exception as e:
        logger.error(f"Error in search_company_info: {str(e)}")
        return f"Error searching: {str(e)}"

# ============================================================================
# QUICK SEARCH FUNCTION - FIXED
# ============================================================================
def quick_search(company_name: str) -> QuickCompanyResearch:
    """
    Quick Search Mode - Generates 6-7 page equivalent report
    Focus: Essential business intelligence with standard depth
    """
    logger.info(f"🚀 Starting QUICK SEARCH for: {company_name}")
    
    try:
        # Gather basic info
        basic_query = f"{company_name} company overview headquarters employees website contact"
        basic_results = search_company_info.invoke({"query": basic_query, "mode": "quick"})
        
        # Gather financial info
        financial_query = f"{company_name} revenue earnings financial results annual report"
        financial_results = search_company_info.invoke({"query": financial_query, "mode": "quick"})
        
        # Gather news
        news_query = f"{company_name} news recent developments leadership changes"
        news_results = search_company_info.invoke({"query": news_query, "mode": "quick"})
        
        # Gather SAP jobs
        job_query = f"{company_name} SAP jobs ERP consultant careers"
        job_results = search_company_info.invoke({"query": job_query, "mode": "quick"})
        
        # Gather SWOT
        swot_query = f"{company_name} competitors strengths weaknesses market position"
        swot_results = search_company_info.invoke({"query": swot_query, "mode": "quick"})
        
        # Compile all results
        all_results = {
            'basic_info': basic_results,
            'financial_info': financial_results,
            'news_updates': news_results,
            'job_postings': job_results,
            'swot_analysis': swot_results
        }
        
        research_context = json.dumps(all_results, indent=2)
        
        system_prompt = f"""
        You are a professional business analyst creating a QUICK SEARCH report (6-7 pages equivalent) for {company_name}.
        
        REPORT SCOPE: Standard business intelligence with essential insights.
        
        Generate a JSON response with ALL fields. For each field:
        - Provide 2-3 sentences of focused, actionable information
        - Be concise but informative
        - For SWOT fields, provide 2-3 key points separated by semicolons
        
        Required JSON structure:
        {{
            "company": "Full company name",
            "employee_count": "Employee count with brief context",
            "annual_revenue": "Revenue figures with growth trend",
            "sic_codes": "SIC code with description",
            "naics_codes": "NAICS code with description",
            "company_official_website": "URL",
            "address": "Full headquarters address",
            "phone_number": "Contact number",
            "recent_funding": "Funding details if available",
            "current_erp": "ERP system in use",
            "leadership_changes": "Recent leadership updates",
            "recent_news": "Top 2-3 recent developments",
            "recent_sap_job_postings": "SAP hiring activity summary",
            "strengths": "2-3 key strengths; separated by semicolons",
            "weaknesses": "2-3 key weaknesses; separated by semicolons",
            "opportunities": "2-3 key opportunities; separated by semicolons",
            "threats": "2-3 key threats; separated by semicolons",
            "founded_year": "Year",
            "ceo_name": "CEO name",
            "industry_sector": "Primary sector",
            "stock_symbol": "Ticker or Private",
            "market_cap": "Market cap if public",
            "headquarters_location": "City, country",
            "business_model": "Brief business model description",
            "competitors": "Top 2-3 competitors with brief comparison",
            "research_timestamp": "{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "search_sources_used": "Multiple search engines",
            "data_freshness": "Current data",
            "search_mode": "Quick Search"
        }}
        
        Return ONLY valid JSON.
        """
        
        human_prompt = f"Extract information for {company_name}:\n\n{research_context}"
        
        messages = [SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)]
        response = llm.invoke(messages)
        
        content = response.content.strip()
        if content.startswith('```json'):
            content = content[7:]
        if content.endswith('```'):
            content = content[:-3]
        content = content.strip()
        
        research_json = json.loads(content)
        research_json['search_mode'] = "Quick Search"
        
        return QuickCompanyResearch(**research_json)
        
    except Exception as e:
        logger.error(f"Quick search error: {str(e)}")
        return QuickCompanyResearch(
            company=company_name,
            employee_count="Error retrieving data",
            annual_revenue="Error retrieving data",
            sic_codes="Error retrieving data",
            naics_codes="Error retrieving data",
            company_official_website="Error retrieving data",
            address="Error retrieving data",
            phone_number="Error retrieving data",
            recent_funding="Error retrieving data",
            current_erp="Error retrieving data",
            leadership_changes="Error retrieving data",
            recent_news="Error retrieving data",
            recent_sap_job_postings="Error retrieving data",
            strengths="Error retrieving data",
            weaknesses="Error retrieving data",
            opportunities="Error retrieving data",
            threats="Error retrieving data",
            search_mode="Quick Search"
        )

# ============================================================================
# DEEP SEARCH FUNCTION - FIXED FOR CONTEXT LENGTH
# ============================================================================
def deep_search(company_name: str) -> DeepCompanyResearch:
    """
    Deep Search Mode - Generates 15-20 page equivalent report
    Focus: Exhaustive analysis with detailed insights and citations
    """
    logger.info(f"🔬 Starting DEEP SEARCH for: {company_name}")
    
    try:
        # Extended search queries for deep mode
        queries = {
            'basic_info': f"{company_name} company overview history founding headquarters employees",
            'financial_deep': f"{company_name} detailed financial analysis revenue breakdown quarterly results investor relations",
            'products': f"{company_name} product portfolio services offerings technology solutions",
            'market': f"{company_name} market position industry analysis competitive landscape market share",
            'technology': f"{company_name} technology stack digital infrastructure IT systems cloud",
            'partnerships': f"{company_name} partnerships alliances strategic relationships joint ventures",
            'leadership': f"{company_name} executive team board directors management leadership",
            'news_deep': f"{company_name} recent news press releases announcements developments",
            'jobs_deep': f"{company_name} SAP job postings ERP careers hiring trends recruitment",
            'swot_deep': f"{company_name} competitive analysis strengths weaknesses opportunities threats",
            'innovation': f"{company_name} research development innovation initiatives patents",
            'sustainability': f"{company_name} sustainability ESG initiatives corporate responsibility",
            'risk': f"{company_name} risk factors challenges regulatory compliance",
            'growth': f"{company_name} growth strategy expansion plans future outlook",
            'customers': f"{company_name} customer segments target market client base"
        }
        
        # Collect results with summarization to reduce token count
        all_results = {}
        all_citations = []
        
        for key, query in queries.items():
            result = search_company_info.invoke({"query": query, "mode": "deep"})
            
            # Extract and store only essential information
            try:
                parsed = json.loads(result)
                
                # Store only top 3 results per query to reduce tokens
                summarized_results = []
                for item in parsed.get('results', [])[:3]:
                    summarized_results.append({
                        'title': item.get('title', '')[:200],  # Limit title length
                        'description': item.get('description', '')[:500],  # Limit description
                        'url': item.get('url', '')
                    })
                    
                    # Collect citations
                    url = item.get('url', '')
                    title = item.get('title', '')
                    if url and title:
                        all_citations.append(f"• {title[:100]}\n  {url}")
                
                all_results[key] = summarized_results
                
            except Exception as e:
                logger.warning(f"Could not parse results from {key}: {str(e)}")
                all_results[key] = []
                continue
        
        # Create condensed research context
        research_context = json.dumps(all_results, indent=1)  # Reduced indentation
        
        # Prepare citations
        if all_citations:
            unique_citations = list(dict.fromkeys(all_citations))[:20]
            citations_text = "\n\n".join(unique_citations)
        else:
            citations_text = "Multiple authoritative sources consulted including business databases, news outlets, official company filings, and professional networks."
        
        # OPTIMIZED SYSTEM PROMPT - More concise but still comprehensive
        system_prompt = f"""You are a senior business intelligence analyst creating a DEEP SEARCH report for {company_name}.

CRITICAL REQUIREMENT: This report MUST generate 15-20 pages of content when converted to a Word document.

Generate a JSON response with ALL fields. For EVERY field you MUST provide:
- MINIMUM 8-12 sentences of comprehensive, detailed analysis
- Include multiple specific examples, dates, figures, and statistics
- Provide historical context and future projections
- Show detailed trends, patterns, and strategic implications
- Add industry comparisons and benchmarking data
- For SWOT fields: MINIMUM 6-8 detailed points, each 3-4 sentences long, separated by semicolons
- For analysis fields: Provide multi-paragraph depth with subsections

WRITING STYLE FOR DEPTH:
- Use detailed explanations with supporting evidence
- Include specific metrics, percentages, and dollar amounts
- Reference multiple time periods (historical, current, projected)
- Compare against industry standards and competitors
- Discuss implications and strategic recommendations
- Add context about market conditions and external factors

Return ONLY valid JSON with this structure:
{{
    "company": "Full official company name",
    "employee_count": "Detailed workforce breakdown with trends",
    "annual_revenue": "Comprehensive financial overview with segments",
    "sic_codes": "SIC classification with explanation",
    "naics_codes": "NAICS codes with explanation",
    "company_official_website": "URL",
    "address": "Full headquarters address",
    "phone_number": "Contact number",
    "recent_funding": "Detailed funding history",
    "current_erp": "Comprehensive ERP analysis",
    "leadership_changes": "Detailed leadership analysis",
    "board_members": "Board composition with backgrounds",
    "recent_news": "Exhaustive news analysis",
    "recent_sap_job_postings": "Detailed SAP hiring analysis",
    "strengths": "4-6 detailed advantages; separated by semicolons",
    "weaknesses": "4-6 detailed vulnerabilities; separated by semicolons",
    "opportunities": "4-6 detailed opportunities; separated by semicolons",
    "threats": "4-6 detailed threats; separated by semicolons",
    "founded_year": "Year with context",
    "ceo_name": "CEO with background",
    "industry_sector": "Detailed sector analysis",
    "stock_symbol": "Ticker or Private",
    "market_cap": "Detailed valuation",
    "headquarters_location": "Complete location",
    "business_model": "Comprehensive business model",
    "competitors": "Detailed competitive analysis of 4-5 competitors",
    "detailed_financials": "Deep financial analysis",
    "product_portfolio": "Complete product/service catalog",
    "customer_segments": "Detailed customer analysis",
    "market_position_analysis": "Comprehensive market position",
    "technology_stack": "Technology infrastructure analysis",
    "partnerships_alliances": "Strategic partnerships",
    "regulatory_compliance": "Regulatory landscape",
    "sustainability_initiatives": "ESG analysis",
    "innovation_rd": "Innovation analysis",
    "geographic_presence": "Global footprint",
    "competitive_advantages": "Sustainable advantages",
    "risk_factors": "Comprehensive risk analysis",
    "growth_strategy": "Strategic growth plan",
    "digital_transformation": "Digital transformation journey",
    "talent_acquisition": "Talent strategy",
    "sources_citations": "4-6 Key sources with proper formatting",
    "research_timestamp": "{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    "search_sources_used": "Multiple premium search engines",
    "data_freshness": "Comprehensive current data",
    "search_mode": "Deep Search"
}}"""
        
        # SHORTENED HUMAN PROMPT - Only essential context
        human_prompt = f"""Extract comprehensive information for {company_name}.

Research Data Summary:
{research_context}

Key Sources:
{citations_text[:2000]}

Provide detailed analysis for all fields."""
        
        messages = [
            SystemMessage(content=system_prompt), 
            HumanMessage(content=human_prompt)
        ]
        
        # Add token count logging for debugging
        logger.info(f"System prompt length: ~{len(system_prompt)} chars")
        logger.info(f"Human prompt length: ~{len(human_prompt)} chars")
        logger.info(f"Total estimated tokens: ~{(len(system_prompt) + len(human_prompt)) / 4}")
        
        response = llm.invoke(messages)
        
        content = response.content.strip()
        if content.startswith('```json'):
            content = content[7:]
        if content.endswith('```'):
            content = content[:-3]
        content = content.strip()
        
        research_json = json.loads(content)
        research_json['search_mode'] = "Deep Search"
        
        # Ensure citations are included
        if 'sources_citations' not in research_json or not research_json['sources_citations']:
            research_json['sources_citations'] = citations_text
        
        return DeepCompanyResearch(**research_json)
        
    except Exception as e:
        logger.error(f"Deep search error: {str(e)}")
        default_citations = "Multiple authoritative sources consulted including business databases, news outlets, official company filings, and professional networks."
        
        return DeepCompanyResearch(
            company=company_name,
            employee_count="Error retrieving data",
            annual_revenue="Error retrieving data",
            sic_codes="Error retrieving data",
            naics_codes="Error retrieving data",
            company_official_website="Error retrieving data",
            address="Error retrieving data",
            phone_number="Error retrieving data",
            recent_funding="Error retrieving data",
            current_erp="Error retrieving data",
            leadership_changes="Error retrieving data",
            board_members="Error retrieving data",
            recent_news="Error retrieving data",
            recent_sap_job_postings="Error retrieving data",
            strengths="Error retrieving data",
            weaknesses="Error retrieving data",
            opportunities="Error retrieving data",
            threats="Error retrieving data",
            sources_citations=default_citations,
            search_mode="Deep Search"
        )


# ============================================================================
# DOCUMENT GENERATION FOR QUICK SEARCH - FIXED WITH CONTENT
# ============================================================================
def create_quick_word_document(research_data: QuickCompanyResearch) -> io.BytesIO:
    """Create a 6-7 page Word document for Quick Search"""
    doc = Document()
    
    def add_formatted_paragraph(doc, label, value):
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        
        # Handle multiline content
        value_str = str(value)
        if ';' in value_str:
            # Split by semicolons and add as separate lines
            lines = value_str.split(';')
            p.add_run(lines[0].strip())
            for line in lines[1:]:
                if line.strip():
                    doc.add_paragraph(f'• {line.strip()}', style='List Bullet')
        else:
            p.add_run(value_str)
    
    # Title
    title = doc.add_heading(f'Quick Search Report: {research_data.company}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    timestamp_para = doc.add_paragraph(f'Research Date: {research_data.research_timestamp}')
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    mode_para = doc.add_paragraph(f'Search Mode: {research_data.search_mode}')
    mode_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mode_para.runs[0]
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # Disclaimer
    note_para = doc.add_paragraph(
        "Important Note: This is AI-generated content. No warranty is made as to its accuracy, "
        "completeness, or suitability for any specific purpose."
    )
    run = note_para.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    # doc.add_paragraph(f'This Quick Search report provides essential business intelligence on {research_data.company}. '
    #                  'The analysis covers fundamental company information, financial overview, and strategic insights.')
    # doc.add_paragraph(f'This report provides a concise overview of {research_data.company}, covering essential company information, '
    #                  'key financial metrics, market positioning, and strategic highlights to support rapid business intelligence needs.')
    doc.add_paragraph(f'This report provides a concise overview of {research_data.company}, covering essential company information, '
                     'key financial metrics, market positioning, and strategic highlights to support rapid business intelligence needs. '
                     f'The analysis examines {research_data.company}\'s core business operations, revenue streams, and competitive advantages '
                     'in its industry. Key financial indicators and recent performance trends are highlighted to provide stakeholders with '
                     'a clear understanding of the company\'s current standing. This summary is designed for quick reference and preliminary '
                     'assessment purposes.')

    # Company Fundamentals
    doc.add_heading('Company Fundamentals', level=1)
    fundamentals = [
        ('Company Name', research_data.company),
        ('Founded', research_data.founded_year),
        ('CEO', research_data.ceo_name),
        ('Industry', research_data.industry_sector),
        ('Headquarters', research_data.headquarters_location),
        ('Employee Count', research_data.employee_count),
        ('Stock Symbol', research_data.stock_symbol)
    ]
    for label, value in fundamentals:
        add_formatted_paragraph(doc, label, value)
    
    # Financial Overview
    doc.add_heading('Financial Overview', level=1)
    financial = [
        ('Annual Revenue', research_data.annual_revenue),
        ('Market Cap', research_data.market_cap),
        ('Recent Funding', research_data.recent_funding)
    ]
    for label, value in financial:
        add_formatted_paragraph(doc, label, value)
    
    # Business Information
    doc.add_heading('Business Information', level=1)
    business = [
        ('Business Model', research_data.business_model),
        ('SIC Codes', research_data.sic_codes),
        ('NAICS Codes', research_data.naics_codes),
        ('Current ERP', research_data.current_erp)
    ]
    for label, value in business:
        add_formatted_paragraph(doc, label, value)
    
    # Recent Developments
    doc.add_heading('Recent Developments', level=1)
    add_formatted_paragraph(doc, 'Recent News', research_data.recent_news)
    doc.add_paragraph()
    add_formatted_paragraph(doc, 'Leadership Changes', research_data.leadership_changes)
    doc.add_paragraph()
    add_formatted_paragraph(doc, 'SAP Job Postings', research_data.recent_sap_job_postings)
    
    # SWOT Analysis
    doc.add_heading('SWOT Analysis', level=1)
    swot = [
        ('Strengths', research_data.strengths),
        ('Weaknesses', research_data.weaknesses),
        ('Opportunities', research_data.opportunities),
        ('Threats', research_data.threats)
    ]
    for category, content in swot:
        doc.add_heading(category, level=2)
        items = [item.strip() for item in str(content).split(';') if item.strip()]
        for item in items:
            if item:
                doc.add_paragraph(item, style='List Bullet')
    
    # Competitive Landscape
    doc.add_heading('Competitive Landscape', level=1)
    add_formatted_paragraph(doc, 'Main Competitors', research_data.competitors)
    
    # Contact Information
    doc.add_heading('Contact Information', level=1)
    contact = [
        ('Website', research_data.company_official_website),
        ('Phone', research_data.phone_number),
        ('Address', research_data.address)
    ]
    for label, value in contact:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))
    
    # Data Quality
    doc.add_heading('Data Quality & Methodology', level=1)
    methodology_text = (
        f"Search Sources: {research_data.search_sources_used}\n"
        f"Data Collection: {research_data.research_timestamp}\n"
        f"Information Freshness: {research_data.data_freshness}\n\n"
        "This Quick Search report uses multi-tool search technology for data collection. "
        "Information is cross-referenced across multiple sources for accuracy."
    )
    doc.add_paragraph(methodology_text)
    
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


# ============================================================================
# DOCUMENT GENERATION FOR DEEP SEARCH - FIXED WITH CONTENT
# ============================================================================
def create_deep_word_document(research_data: DeepCompanyResearch) -> io.BytesIO:
    """Create a 15-20 page Word document for Deep Search with citations and TOC"""
    doc = Document()
    
    def add_formatted_paragraph(doc, label, value):
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        
        # Handle multiline content
        value_str = str(value)
        if ';' in value_str:
            # Split by semicolons and add as separate lines
            lines = value_str.split(';')
            p.add_run(lines[0].strip())
            for line in lines[1:]:
                if line.strip():
                    doc.add_paragraph(f'• {line.strip()}', style='List Bullet')
        else:
            p.add_run(value_str)
    
    # Title
    title = doc.add_heading(f'Deep Search Report: {research_data.company}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    timestamp_para = doc.add_paragraph(f'Research Date: {research_data.research_timestamp}')
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    mode_para = doc.add_paragraph(f'Search Mode: {research_data.search_mode} (Comprehensive Analysis)')
    mode_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mode_para.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 255)
    
    doc.add_page_break()
    
    # Disclaimer
    note_para = doc.add_paragraph(
        "Important Note: This is AI-generated comprehensive research content. "
        "No warranty is made as to its accuracy, completeness, or suitability for any specific purpose. "
        "This report includes extensive analysis and should be verified for critical business decisions."
    )
    run = note_para.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    # doc.add_paragraph(f'This Deep Search report provides exhaustive business intelligence and strategic analysis on {research_data.company}. '
    #                  'The comprehensive analysis covers fundamental company information, detailed financial analysis, competitive positioning, '
    #                  'strategic initiatives, risk factors, and growth opportunities. This report is designed for strategic decision-making '
    #                  'and includes citations to source materials.')
    # doc.add_paragraph(f'This comprehensive report provides an in-depth analysis of {research_data.company}, examining its business model, '
    #                  'market position, financial performance, competitive landscape, and strategic direction. The report synthesizes '
    #                  'information from multiple sources to deliver actionable intelligence for strategic decision-making, investment analysis, '
    #                  'and competitive assessment.')
    doc.add_paragraph(f'This comprehensive report provides an in-depth analysis of {research_data.company}, examining its business model, '
                     'market position, financial performance, competitive landscape, and strategic direction. The report synthesizes '
                     'information from multiple sources to deliver actionable intelligence for strategic decision-making, investment analysis, '
                     'and competitive assessment. '
                     f'The analysis delves into {research_data.company}\'s organizational structure, product and service offerings, target markets, '
                     'and value proposition within its industry sector. Detailed financial metrics, including revenue trends, profitability indicators, '
                     'and key performance ratios, are examined to assess the company\'s fiscal health and operational efficiency. '
                     'The competitive analysis identifies major rivals, market dynamics, and differentiation strategies that position '
                     f'{research_data.company} within its competitive ecosystem. Strategic initiatives, growth opportunities, and potential '
                     'risk factors are evaluated to provide a forward-looking perspective. This report serves as a comprehensive resource '
                     'for executives, investors, and analysts requiring detailed business intelligence.')

    
    # Table of Contents (IN WORD DOC ONLY)
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Company Fundamentals',
        '2. Comprehensive Financial Analysis',
        '3. Business Operations & Strategy',
        '4. Leadership & Governance',
        '5. Strategic Developments & Initiatives',
        '6. Technology & SAP Landscape',
        '7. Market Position & Competitive Analysis',
        '8. Comprehensive SWOT Analysis',
        '9. Risk Assessment & Compliance',
        '10. Growth Strategy & Future Outlook',
        '11. Contact Information',
        '12. Source Citations'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Company Fundamentals (Extended)
    doc.add_heading('1. Company Fundamentals', level=1)
    fundamentals = [
        ('Company Name', research_data.company),
        ('Founded', research_data.founded_year),
        ('CEO & Leadership', research_data.ceo_name),
        ('Industry Sector', research_data.industry_sector),
        ('Headquarters', research_data.headquarters_location),
        ('Geographic Presence', research_data.geographic_presence),
        ('Employee Count', research_data.employee_count),
        ('Stock Symbol', research_data.stock_symbol),
        ('Business Model', research_data.business_model)
    ]
    for label, value in fundamentals:
        add_formatted_paragraph(doc, label, value)
        doc.add_paragraph()
    
    # Financial Analysis (Deep)
    doc.add_heading('2. Comprehensive Financial Analysis', level=1)
    financial = [
        ('Annual Revenue', research_data.annual_revenue),
        ('Market Capitalization', research_data.market_cap),
        ('Detailed Financials', research_data.detailed_financials),
        ('Recent Funding', research_data.recent_funding)
    ]
    for label, value in financial:
        add_formatted_paragraph(doc, label, value)
        doc.add_paragraph()
    
    # Business Operations
    doc.add_heading('3. Business Operations & Strategy', level=1)
    operations = [
        ('Product Portfolio', research_data.product_portfolio),
        ('Customer Segments', research_data.customer_segments),
        ('SIC Codes', research_data.sic_codes),
        ('NAICS Codes', research_data.naics_codes),
        ('Current ERP System', research_data.current_erp),
        ('Technology Stack', research_data.technology_stack)
    ]
    for label, value in operations:
        add_formatted_paragraph(doc, label, value)
        doc.add_paragraph()
    
    # Leadership & Governance
    doc.add_heading('4. Leadership & Governance', level=1)
    leadership = [
        ('Chief Executive Officer', research_data.ceo_name),
        ('Board of Directors', research_data.board_members),
        ('Recent Leadership Changes', research_data.leadership_changes),
        ('Talent Acquisition Strategy', research_data.talent_acquisition)
    ]
    for label, value in leadership:
        add_formatted_paragraph(doc, label, value)
        doc.add_paragraph()
    
    # Strategic Developments
    doc.add_heading('5. Strategic Developments & Initiatives', level=1)
    strategic = [
        ('Recent News & Announcements', research_data.recent_news),
        ('Partnerships & Alliances', research_data.partnerships_alliances),
        ('Digital Transformation', research_data.digital_transformation),
        ('Innovation & R&D', research_data.innovation_rd),
        ('Sustainability Initiatives', research_data.sustainability_initiatives)
    ]
    for label, value in strategic:
        add_formatted_paragraph(doc, label, value)
        doc.add_paragraph()
    
    # Technology & SAP
    doc.add_heading('6. Technology & SAP Landscape', level=1)
    add_formatted_paragraph(doc, 'SAP Job Postings Analysis', research_data.recent_sap_job_postings)
    doc.add_paragraph()
    add_formatted_paragraph(doc, 'Technology Infrastructure', research_data.technology_stack)
    doc.add_paragraph()
    
    # Market Position
    doc.add_heading('7. Market Position & Competitive Analysis', level=1)
    market = [
        ('Market Position Analysis', research_data.market_position_analysis),
        ('Competitive Advantages', research_data.competitive_advantages),
        ('Main Competitors', research_data.competitors)
    ]
    for label, value in market:
        add_formatted_paragraph(doc, label, value)
        doc.add_paragraph()
    
    # SWOT Analysis (Deep)
    doc.add_heading('8. Comprehensive SWOT Analysis', level=1)
    swot = [
        ('Strengths', research_data.strengths),
        ('Weaknesses', research_data.weaknesses),
        ('Opportunities', research_data.opportunities),
        ('Threats', research_data.threats)
    ]
    for category, content in swot:
        doc.add_heading(category, level=2)
        items = [item.strip() for item in str(content).split(';') if item.strip()]
        for item in items:
            if item:
                doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()
    
    # Risk Assessment
    doc.add_heading('9. Risk Assessment & Compliance', level=1)
    risk = [
        ('Risk Factors', research_data.risk_factors),
        ('Regulatory Compliance', research_data.regulatory_compliance)
    ]
    for label, value in risk:
        add_formatted_paragraph(doc, label, value)
        doc.add_paragraph()
    
    # Growth Strategy
    doc.add_heading('10. Growth Strategy & Future Outlook', level=1)
    add_formatted_paragraph(doc, 'Strategic Growth Plan', research_data.growth_strategy)
    doc.add_paragraph()
    
    # Contact Information
    doc.add_heading('11. Contact Information', level=1)
    contact = [
        ('Website', research_data.company_official_website),
        ('Phone', research_data.phone_number),
        ('Address', research_data.address)
    ]
    for label, value in contact:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))
    
    doc.add_paragraph()
    
    # Citations Section
    doc.add_heading('12. Source Citations', level=1)
    doc.add_paragraph('This comprehensive report was compiled using information from multiple authoritative sources. '
                     'Key sources consulted include:')
    doc.add_paragraph()
    
    # Add citations
    citations_text = str(research_data.sources_citations)
    citations = citations_text.split('\n\n')
    for i, citation in enumerate(citations[:20], 1):
        if citation.strip():
            doc.add_paragraph(f'{i}. {citation.strip()}', style='List Number')
    
    # Methodology
    doc.add_heading('Research Methodology', level=1)
    methodology_text = (
        f"Search Engines Used: {research_data.search_sources_used}\n\n"
        f"Data Collection: {research_data.research_timestamp}\n\n"
        f"Information Freshness: {research_data.data_freshness}\n\n"
        "This Deep Search report employs advanced multi-tool search technology with exhaustive data collection "
        "across multiple premium sources. Information is cross-referenced extensively for accuracy validation. "
        "The comprehensive nature of this report makes it suitable for strategic planning, investment analysis, "
        "and detailed competitive intelligence."
    )
    doc.add_paragraph(methodology_text)
    
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# HTML Report Generation
def generate_quick_html_report(report_dict):
    """Generate HTML for Quick Search (6-7 pages)"""
    return f"""
    <h3>Quick Search Report - Essential Business Intelligence</h3>
    <div style="background: #e3f2fd; color: black; padding: 15px; border-radius: 10px; margin-bottom: 20px; border-left: 4px solid #2196F3;">
        <div style="font-weight: bold; font-size: 16px; margin-bottom: 10px;">🚀 Quick Search Mode</div>
        <div style="display: flex; gap: 20px; flex-wrap: wrap;">
            <div><strong>Report Type:</strong> {report_dict.get('search_mode', 'Consice Summary')}</div>
            <div><strong>Generated:</strong> {report_dict.get('research_timestamp', 'N/A')}</div>
        </div>
    </div>

    <h3>Company Overview</h3>
    <ul>
        <li><strong>Company:</strong> {report_dict['company']}</li>
        <li><strong>Founded:</strong> {report_dict['founded_year']}</li>
        <li><strong>CEO:</strong> {report_dict['ceo_name']}</li>
        <li><strong>Industry:</strong> {report_dict['industry_sector']}</li>
        <li><strong>Headquarters:</strong> {report_dict['headquarters_location']}</li>
        <li><strong>Employees:</strong> {report_dict['employee_count']}</li>
    </ul>

    <h3>Financial Snapshot</h3>
    <ul>
        <li><strong>Revenue:</strong> {report_dict['annual_revenue']}</li>
        <li><strong>Market Cap:</strong> {report_dict['market_cap']}</li>
        <li><strong>Stock Symbol:</strong> {report_dict['stock_symbol']}</li>
        <li><strong>Funding:</strong> {report_dict['recent_funding']}</li>
    </ul>

    <h3>Business Details</h3>
    <ul>
        <li><strong>Business Model:</strong> {report_dict['business_model']}</li>
        <li><strong>SIC Codes:</strong> {report_dict['sic_codes']}</li>
        <li><strong>NAICS Codes:</strong> {report_dict['naics_codes']}</li>
        <li><strong>Current ERP:</strong> {report_dict['current_erp']}</li>
    </ul>

    <h3>Recent Developments</h3>
    <ul>
        <li><strong>Latest News:</strong> {report_dict['recent_news']}</li>
        <li><strong>Leadership Changes:</strong> {report_dict['leadership_changes']}</li>
        <li><strong>SAP Hiring:</strong> {report_dict['recent_sap_job_postings']}</li>
    </ul>

    <h3>SWOT Analysis</h3>
    <ul>
        <li><strong>Strengths:</strong> {report_dict['strengths']}</li>
        <li><strong>Weaknesses:</strong> {report_dict['weaknesses']}</li>
        <li><strong>Opportunities:</strong> {report_dict['opportunities']}</li>
        <li><strong>Threats:</strong> {report_dict['threats']}</li>
    </ul>

    <h3>Competitors</h3>
    <ul>
        <li><strong>Main Competitors:</strong> {report_dict['competitors']}</li>
    </ul>

    <h3>Contact</h3>
    <ul>
        <li><strong>Website:</strong> {report_dict['company_official_website']}</li>
        <li><strong>Phone:</strong> {report_dict['phone_number']}</li>
        <li><strong>Address:</strong> {report_dict['address']}</li>
    </ul>
    """

def generate_deep_html_report(report_dict):
    """Generate HTML for Deep Search (15-20 pages)"""
    return f"""
    <h3>Deep Search Report - Comprehensive Strategic Analysis</h3>
    <div style="background: #f3e5f5; color: black; padding: 15px; border-radius: 10px; margin-bottom: 20px; border-left: 4px solid #9C27B0;">
        <div style="font-weight: bold; font-size: 16px; margin-bottom: 10px;">🔬 Deep Search Mode</div>
        <div style="display: flex; gap: 20px; flex-wrap: wrap;">
            <div><strong>Report Type:</strong> {report_dict.get('search_mode', 'Comprehensive Summary')}</div>
            <div><strong>Generated:</strong> {report_dict.get('research_timestamp', 'N/A')}</div>
        </div>
    </div>

    <h2>📋 Table of Contents</h2>
    <ul style="background: #fafafa; padding: 15px; border-radius: 5px;">
        <li>1. Company Fundamentals</li>
        <li>2. Comprehensive Financial Analysis</li>
        <li>3. Business Operations & Strategy</li>
        <li>4. Leadership & Governance</li>
        <li>5. Strategic Developments</li>
        <li>6. Technology & SAP Landscape</li>
        <li>7. Market Position & Competition</li>
        <li>8. SWOT Analysis</li>
        <li>9. Risk Assessment</li>
        <li>10. Growth Strategy</li>
        <li>11. Source Citations</li>
    </ul>

    <h2>1. Company Fundamentals</h2>
    <ul>
        <li><strong>Company:</strong> {report_dict['company']}</li>
        <li><strong>Founded:</strong> {report_dict['founded_year']}</li>
        <li><strong>CEO:</strong> {report_dict['ceo_name']}</li>
        <li><strong>Industry:</strong> {report_dict['industry_sector']}</li>
        <li><strong>Headquarters:</strong> {report_dict['headquarters_location']}</li>
        <li><strong>Geographic Presence:</strong> {report_dict.get('geographic_presence', 'Not found')}</li>
        <li><strong>Workforce Analysis:</strong> {report_dict['employee_count']}</li>
        <li><strong>Business Model:</strong> {report_dict['business_model']}</li>
    </ul>

    <h2>2. Comprehensive Financial Analysis</h2>
    <ul>
        <li><strong>Revenue Analysis:</strong> {report_dict['annual_revenue']}</li>
        <li><strong>Market Capitalization:</strong> {report_dict['market_cap']}</li>
        <li><strong>Stock Performance:</strong> {report_dict['stock_symbol']}</li>
        <li><strong>Detailed Financials:</strong> {report_dict.get('detailed_financials', 'Not found')}</li>
        <li><strong>Funding & Investments:</strong> {report_dict['recent_funding']}</li>
    </ul>

    <h2>3. Business Operations & Strategy</h2>
    <ul>
        <li><strong>Product Portfolio:</strong> {report_dict.get('product_portfolio', 'Not found')}</li>
        <li><strong>Customer Segments:</strong> {report_dict.get('customer_segments', 'Not found')}</li>
        <li><strong>SIC Classification:</strong> {report_dict['sic_codes']}</li>
        <li><strong>NAICS Classification:</strong> {report_dict['naics_codes']}</li>
        <li><strong>ERP System:</strong> {report_dict['current_erp']}</li>
        <li><strong>Technology Stack:</strong> {report_dict.get('technology_stack', 'Not found')}</li>
    </ul>

    <h2>4. Leadership & Governance</h2>
    <ul>
        <li><strong>Chief Executive:</strong> {report_dict['ceo_name']}</li>
        <li><strong>Board Composition:</strong> {report_dict.get('board_members', 'Not found')}</li>
        <li><strong>Leadership Evolution:</strong> {report_dict['leadership_changes']}</li>
        <li><strong>Talent Strategy:</strong> {report_dict.get('talent_acquisition', 'Not found')}</li>
    </ul>

    <h2>5. Strategic Developments & Initiatives</h2>
    <ul>
        <li><strong>Recent Developments:</strong> {report_dict['recent_news']}</li>
        <li><strong>Strategic Partnerships:</strong> {report_dict.get('partnerships_alliances', 'Not found')}</li>
        <li><strong>Digital Transformation:</strong> {report_dict.get('digital_transformation', 'Not found')}</li>
        <li><strong>Innovation & R&D:</strong> {report_dict.get('innovation_rd', 'Not found')}</li>
        <li><strong>Sustainability Programs:</strong> {report_dict.get('sustainability_initiatives', 'Not found')}</li>
    </ul>

    <h2>6. Technology & SAP Landscape</h2>
    <ul>
        <li><strong>SAP Hiring Analysis:</strong> {report_dict['recent_sap_job_postings']}</li>
        <li><strong>Technology Infrastructure:</strong> {report_dict.get('technology_stack', 'Not found')}</li>
        <li><strong>Digital Maturity:</strong> {report_dict.get('digital_transformation', 'Not found')}</li>
    </ul>

    <h2>7. Market Position & Competitive Analysis</h2>
    <ul>
        <li><strong>Market Position:</strong> {report_dict.get('market_position_analysis', 'Not found')}</li>
        <li><strong>Competitive Advantages:</strong> {report_dict.get('competitive_advantages', 'Not found')}</li>
        <li><strong>Competitive Landscape:</strong> {report_dict['competitors']}</li>
    </ul>

    <h2>8. Comprehensive SWOT Analysis</h2>
    <h3 style="color: #2e7d32;">Strengths</h3>
    <ul>
        <li>{report_dict['strengths'].replace(';', '</li><li>')}</li>
    </ul>
    <h3 style="color: #c62828;">Weaknesses</h3>
    <ul>
        <li>{report_dict['weaknesses'].replace(';', '</li><li>')}</li>
    </ul>
    <h3 style="color: #1565c0;">Opportunities</h3>
    <ul>
        <li>{report_dict['opportunities'].replace(';', '</li><li>')}</li>
    </ul>
    <h3 style="color: #e65100;">Threats</h3>
    <ul>
        <li>{report_dict['threats'].replace(';', '</li><li>')}</li>
    </ul>

    <h2>9. Risk Assessment & Compliance</h2>
    <ul>
        <li><strong>Risk Factors:</strong> {report_dict.get('risk_factors', 'Not found')}</li>
        <li><strong>Regulatory Compliance:</strong> {report_dict.get('regulatory_compliance', 'Not found')}</li>
    </ul>

    <h2>10. Growth Strategy & Future Outlook</h2>
    <ul>
        <li><strong>Strategic Growth Plan:</strong> {report_dict.get('growth_strategy', 'Not found')}</li>
    </ul>

    <h2>11. Contact Information</h2>
    <ul>
        <li><strong>Website:</strong> {report_dict['company_official_website']}</li>
        <li><strong>Phone:</strong> {report_dict['phone_number']}</li>
        <li><strong>Address:</strong> {report_dict['address']}</li>
    </ul>

    <h2>12. Source Citations</h2>
    <div style="background: #f5f5f5; padding: 15px; border-radius: 5px; font-size: 14px;">
        <p><strong>Key Sources Consulted:</strong></p>
        <p style="white-space: pre-line;">{report_dict.get('sources_citations', 'Multiple authoritative sources')}</p>
    </div>

    <div style="background: #fff3e0; padding: 15px; border-radius: 5px; margin-top: 20px; border-left: 4px solid #ff9800;">
        <p><strong>Research Methodology:</strong></p>
        <p>This comprehensive Deep Search report employs exhaustive multi-tool search technology with extensive data collection 
        across multiple premium sources. Information is cross-referenced extensively for accuracy validation. 
        The detailed nature of this 15-20 page equivalent report makes it suitable for strategic planning, 
        investment analysis, and detailed competitive intelligence.</p>
    </div>
    """

# Flask Application
import logging
import uuid
from datetime import datetime
from flask import Flask, render_template, request, send_file, jsonify, session
from flask import Flask, redirect, url_for

app = Flask(__name__)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# In-memory search history { session_id: [history_items] }
search_history = {}

app.secret_key = 'your-secret-key-here-change-this-in-production'

# HTML Report Generation for Quick Search
def generate_quick_html_report(report_dict):
    """Generate HTML for Quick Search (6-7 pages)"""
    search_sources = report_dict.get('search_sources_used', 'Multiple sources')
    data_freshness = report_dict.get('data_freshness', 'Recent data')
    
    return f"""
    <h3>Quick Search Report - Essential Business Intelligence</h3>
    <div style="background: #e3f2fd; color: black; padding: 15px; border-radius: 10px; margin-bottom: 20px; border-left: 4px solid #2196F3;">
        <div style="display: flex; gap: 20px; flex-wrap: wrap;">
            <div><strong>Report Type:</strong> {report_dict.get('search_mode', 'Concise Summary')}</div>
            <div><strong>Generated:</strong> {report_dict.get('research_timestamp', 'N/A')}</div>
            <div><strong>Data Freshness:</strong> {data_freshness}</div>
        </div>
    </div>

    <h3>Company Overview</h3>
    <ul>
        <li><strong>Company:</strong> {report_dict['company']}</li>
        <li><strong>Founded:</strong> {report_dict.get('founded_year', 'Not found')}</li>
        <li><strong>CEO:</strong> {report_dict.get('ceo_name', 'Not found')}</li>
        <li><strong>Industry:</strong> {report_dict.get('industry_sector', 'Not found')}</li>
        <li><strong>Headquarters:</strong> {report_dict.get('headquarters_location', 'Not found')}</li>
        <li><strong>Employees:</strong> {report_dict['employee_count']}</li>
    </ul>

    <h3>Financial Snapshot</h3>
    <ul>
        <li><strong>Revenue:</strong> {report_dict['annual_revenue']}</li>
        <li><strong>Market Cap:</strong> {report_dict.get('market_cap', 'Not found')}</li>
        <li><strong>Stock Symbol:</strong> {report_dict.get('stock_symbol', 'Not found')}</li>
        <li><strong>Funding:</strong> {report_dict['recent_funding']}</li>
    </ul>

    <h3>Business Details</h3>
    <ul>
        <li><strong>Business Model:</strong> {report_dict.get('business_model', 'Not found')}</li>
        <li><strong>SIC Codes:</strong> {report_dict['sic_codes']}</li>
        <li><strong>NAICS Codes:</strong> {report_dict['naics_codes']}</li>
        <li><strong>Current ERP:</strong> {report_dict['current_erp']}</li>
    </ul>

    <h3>Recent Developments</h3>
    <ul>
        <li><strong>Latest News:</strong> {report_dict['recent_news']}</li>
        <li><strong>Leadership Changes:</strong> {report_dict['leadership_changes']}</li>
        <li><strong>SAP Hiring:</strong> {report_dict['recent_sap_job_postings']}</li>
    </ul>

    <h3>SWOT Analysis</h3>
    <ul>
        <li><strong>Strengths:</strong> {report_dict['strengths']}</li>
        <li><strong>Weaknesses:</strong> {report_dict['weaknesses']}</li>
        <li><strong>Opportunities:</strong> {report_dict['opportunities']}</li>
        <li><strong>Threats:</strong> {report_dict['threats']}</li>
    </ul>

    <h3>Competitors</h3>
    <ul>
        <li><strong>Main Competitors:</strong> {report_dict.get('competitors', 'Not found')}</li>
    </ul>

    <h3>Contact</h3>
    <ul>
        <li><strong>Website:</strong> {report_dict['company_official_website']}</li>
        <li><strong>Phone:</strong> {report_dict['phone_number']}</li>
        <li><strong>Address:</strong> {report_dict['address']}</li>
    </ul>

    <h3>Disclaimer</h3>
    <p>This Quick Search report uses AI-based multi-tool search technology. While cross-referenced across multiple sources, 
    some data may be incomplete or subject to change. Please verify through official channels for critical business decisions.</p>
    """

# HTML Report Generation for Deep Search
def generate_deep_html_report(report_dict):
    """Generate HTML for Deep Search (15-20 pages)"""
    search_sources = report_dict.get('search_sources_used', 'Multiple sources')
    data_freshness = report_dict.get('data_freshness', 'Recent data')
    
    return f"""
    <h3>Deep Search Report - Comprehensive Strategic Analysis</h3>
    <div style="background: #e3f2fd; color: black; padding: 15px; border-radius: 10px; margin-bottom: 20px; border-left: 4px solid #9C27B0;">
        <div style="display: flex; gap: 20px; flex-wrap: wrap;">
            <div><strong>Report Type:</strong> {report_dict.get('search_mode', 'Comprehensive Summary')}</div>
            <div><strong>Generated:</strong> {report_dict.get('research_timestamp', 'N/A')}</div>
            <div><strong>Data Freshness:</strong> {data_freshness}</div>
        </div>
    </div>

    <h2>1. Company Fundamentals</h2>
    <ul>
        <li><strong>Company:</strong> {report_dict['company']}</li>
        <li><strong>Founded:</strong> {report_dict.get('founded_year', 'Not found')}</li>
        <li><strong>CEO:</strong> {report_dict.get('ceo_name', 'Not found')}</li>
        <li><strong>Industry:</strong> {report_dict.get('industry_sector', 'Not found')}</li>
        <li><strong>Headquarters:</strong> {report_dict.get('headquarters_location', 'Not found')}</li>
        <li><strong>Geographic Presence:</strong> {report_dict.get('geographic_presence', 'Not found')}</li>
        <li><strong>Workforce Analysis:</strong> {report_dict['employee_count']}</li>
        <li><strong>Business Model:</strong> {report_dict.get('business_model', 'Not found')}</li>
    </ul>

    <h2>2. Comprehensive Financial Analysis</h2>
    <ul>
        <li><strong>Revenue Analysis:</strong> {report_dict['annual_revenue']}</li>
        <li><strong>Market Capitalization:</strong> {report_dict.get('market_cap', 'Not found')}</li>
        <li><strong>Stock Performance:</strong> {report_dict.get('stock_symbol', 'Not found')}</li>
        <li><strong>Detailed Financials:</strong> {report_dict.get('detailed_financials', 'Not found')}</li>
        <li><strong>Funding & Investments:</strong> {report_dict['recent_funding']}</li>
    </ul>

    <h2>3. Business Operations & Strategy</h2>
    <ul>
        <li><strong>Product Portfolio:</strong> {report_dict.get('product_portfolio', 'Not found')}</li>
        <li><strong>Customer Segments:</strong> {report_dict.get('customer_segments', 'Not found')}</li>
        <li><strong>SIC Classification:</strong> {report_dict['sic_codes']}</li>
        <li><strong>NAICS Classification:</strong> {report_dict['naics_codes']}</li>
        <li><strong>ERP System:</strong> {report_dict['current_erp']}</li>
        <li><strong>Technology Stack:</strong> {report_dict.get('technology_stack', 'Not found')}</li>
    </ul>

    <h2>4. Leadership & Governance</h2>
    <ul>
        <li><strong>Chief Executive:</strong> {report_dict.get('ceo_name', 'Not found')}</li>
        <li><strong>Board Composition:</strong> {report_dict.get('board_members', 'Not found')}</li>
        <li><strong>Leadership Evolution:</strong> {report_dict['leadership_changes']}</li>
        <li><strong>Talent Strategy:</strong> {report_dict.get('talent_acquisition', 'Not found')}</li>
    </ul>

    <h2>5. Strategic Developments & Initiatives</h2>
    <ul>
        <li><strong>Recent Developments:</strong> {report_dict['recent_news']}</li>
        <li><strong>Strategic Partnerships:</strong> {report_dict.get('partnerships_alliances', 'Not found')}</li>
        <li><strong>Digital Transformation:</strong> {report_dict.get('digital_transformation', 'Not found')}</li>
        <li><strong>Innovation & R&D:</strong> {report_dict.get('innovation_rd', 'Not found')}</li>
        <li><strong>Sustainability Programs:</strong> {report_dict.get('sustainability_initiatives', 'Not found')}</li>
    </ul>

    <h2>6. Technology & SAP Landscape</h2>
    <ul>
        <li><strong>SAP Hiring Analysis:</strong> {report_dict['recent_sap_job_postings']}</li>
        <li><strong>Technology Infrastructure:</strong> {report_dict.get('technology_stack', 'Not found')}</li>
        <li><strong>Digital Maturity:</strong> {report_dict.get('digital_transformation', 'Not found')}</li>
    </ul>

    <h2>7. Market Position & Competitive Analysis</h2>
    <ul>
        <li><strong>Market Position:</strong> {report_dict.get('market_position_analysis', 'Not found')}</li>
        <li><strong>Competitive Advantages:</strong> {report_dict.get('competitive_advantages', 'Not found')}</li>
        <li><strong>Competitive Landscape:</strong> {report_dict.get('competitors', 'Not found')}</li>
    </ul>

    <h2>8. Comprehensive SWOT Analysis</h2>
    <h3 style="color: #2e7d32;">Strengths</h3>
    <ul>
        <li>{report_dict['strengths'].replace(';', '</li><li>')}</li>
    </ul>
    <h3 style="color: #c62828;">Weaknesses</h3>
    <ul>
        <li>{report_dict['weaknesses'].replace(';', '</li><li>')}</li>
    </ul>
    <h3 style="color: #1565c0;">Opportunities</h3>
    <ul>
        <li>{report_dict['opportunities'].replace(';', '</li><li>')}</li>
    </ul>
    <h3 style="color: #e65100;">Threats</h3>
    <ul>
        <li>{report_dict['threats'].replace(';', '</li><li>')}</li>
    </ul>

    <h2>9. Risk Assessment & Compliance</h2>
    <ul>
        <li><strong>Risk Factors:</strong> {report_dict.get('risk_factors', 'Not found')}</li>
        <li><strong>Regulatory Compliance:</strong> {report_dict.get('regulatory_compliance', 'Not found')}</li>
    </ul>

    <h2>10. Growth Strategy & Future Outlook</h2>
    <ul>
        <li><strong>Strategic Growth Plan:</strong> {report_dict.get('growth_strategy', 'Not found')}</li>
    </ul>

    <h2>11. Contact Information</h2>
    <ul>
        <li><strong>Website:</strong> {report_dict['company_official_website']}</li>
        <li><strong>Phone:</strong> {report_dict['phone_number']}</li>
        <li><strong>Address:</strong> {report_dict['address']}</li>
    </ul>

    <h2>12. Source Citations</h2>
    <div style="background: #f5f5f5; padding: 15px; border-radius: 5px; font-size: 14px;">
        <p><strong>Key Sources Consulted:</strong></p>
        <div style="white-space: pre-line; line-height: 1.8;">
            {report_dict.get('sources_citations', 'Multiple authoritative sources consulted including business databases, news outlets, official filings, and professional networks.')}
        </div>
    </div>

    <div style="background: #fff3e0; padding: 15px; border-radius: 5px; margin-top: 20px; border-left: 4px solid #ff9800;">
        <p><strong>Research Methodology:</strong></p>
        <p>This comprehensive Deep Search report employs exhaustive multi-tool search technology with extensive data collection 
        across multiple premium sources. Information is cross-referenced extensively for accuracy validation. 
        The detailed nature of this 15-20 page equivalent report makes it suitable for strategic planning, 
        investment analysis, and detailed competitive intelligence.</p>
    </div>

    <h3>Disclaimer</h3>
    <p>This Deep Search report uses advanced AI-based multi-tool search technology with comprehensive data collection. 
    While extensively cross-referenced across multiple authoritative sources, some data may be incomplete or subject to change. 
    Please verify through official channels for critical business decisions.</p>
    """

# ---------------- Helper: get session id ---------------- #
def get_session_id():
    """Return a stable session ID for this browser session"""
    if "session_id" not in session:
        session["session_id"] = str(uuid.uuid4())
    return session["session_id"]

# @app.route("/", methods=["GET", "POST"])
# def index():
#     company_name = None
#     error = None
#     generated_report = None
#     filename = None
#     search_mode = "Quick Search"

#     sid = get_session_id()
#     if sid not in search_history:
#         search_history[sid] = []

#     if request.method == "POST":
#         company_name = request.form.get("company", "").strip()
#         search_mode = request.form.get("search_mode", "Quick Search").strip()
        
#         if not company_name:
#             error = "Please enter a company name"
#         else:
#             try:
#                 logger.info(f"🚀 Starting {search_mode} for: {company_name}")

#                 # Execute appropriate search based on mode
#                 if search_mode == "Deep Search":
#                     # Call deep_search function
#                     report_obj = deep_search(company_name)
#                     report_dict = report_obj.model_dump()
#                     generated_report = generate_deep_html_report(report_dict)
#                 else:
#                     # Call quick_search function
#                     report_obj = quick_search(company_name)
#                     report_dict = report_obj.model_dump()
#                     generated_report = generate_quick_html_report(report_dict)

#                 filename = f"{company_name.replace(' ', '_')}_{search_mode.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

#                 # Save into memory history
#                 history_item = {
#                     "company": report_dict["company"],
#                     "timestamp": datetime.now().strftime("%H:%M:%S"),
#                     "search_mode": search_mode,
#                     "report_html": generated_report,
#                     "filename": filename,
#                     "report_dict": report_dict
#                 }
#                 search_history[sid].insert(0, history_item)

#             except Exception as e:
#                 logging.error(f"Error during research: {str(e)}")
#                 error = f"An error occurred: {str(e)}"

#     # For GET (like "New Research"), just show empty report but preserve history
#     return render_template(
#         "test.html",
#         company=company_name,
#         error=error,
#         generated_report=generated_report,
#         file_ready=bool(generated_report),
#         filename=filename,
#         history=search_history[sid],
#     )

@app.route("/", methods=["GET", "POST"])
def index():
    company_name = None
    error = None
    generated_report = None
    filename = None
    search_mode = "Quick Search"

    sid = get_session_id()
    if sid not in search_history:
        search_history[sid] = []

    if request.method == "POST":
        company_name = request.form.get("company", "").strip()
        search_mode = request.form.get("search_mode", "Quick Search").strip()
        
        if not company_name:
            error = "Please enter a company name"
            return render_template(
                "test.html",
                company=company_name,
                error=error,
                generated_report=None,
                file_ready=False,
                filename=None,
                history=search_history[sid],
            )
        else:
            try:
                logger.info(f"🚀 Starting {search_mode} for: {company_name}")

                if search_mode == "Deep Search":
                    report_obj = deep_search(company_name)
                    report_dict = report_obj.model_dump()
                    generated_report = generate_deep_html_report(report_dict)
                else:
                    report_obj = quick_search(company_name)
                    report_dict = report_obj.model_dump()
                    generated_report = generate_quick_html_report(report_dict)

                filename = f"{company_name.replace(' ', '_')}_{search_mode.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

                history_item = {
                    "company": report_dict["company"],
                    "timestamp": datetime.now().strftime("%H:%M:%S"),
                    "search_mode": search_mode,
                    "report_html": generated_report,
                    "filename": filename,
                    "report_dict": report_dict
                }
                search_history[sid].insert(0, history_item)

                # REDIRECT after POST to prevent re-submission
                return redirect(url_for('index'))

            except Exception as e:
                logging.error(f"Error during research: {str(e)}")
                error = f"An error occurred: {str(e)}"
                return render_template(
                    "test.html",
                    company=company_name,
                    error=error,
                    generated_report=None,
                    file_ready=False,
                    filename=None,
                    history=search_history[sid],
                )

    # GET request handling
    # Check if this is a "New Research" request
    new_research = request.args.get('new', 'false') == 'true'
    
    if new_research:
        # Show blank form for new research
        return render_template(
            "test.html",
            company=None,
            error=None,
            generated_report=None,
            file_ready=False,
            filename=None,
            history=search_history[sid],
        )
    else:
        # Show the latest report if available
        latest_report = search_history[sid][0] if search_history[sid] else None
        
        if latest_report:
            company_name = latest_report["company"]
            generated_report = latest_report["report_html"]
            filename = latest_report["filename"]

        return render_template(
            "test.html",
            company=company_name,
            error=error,
            generated_report=generated_report,
            file_ready=bool(generated_report),
            filename=filename,
            history=search_history[sid],
        )

@app.route("/history/<int:idx>")
def load_history(idx):
    """Load a previously generated report from memory"""
    sid = get_session_id()
    if sid not in search_history or idx < 0 or idx >= len(search_history[sid]):
        return render_template("test.html", error="History not found", history=[])

    item = search_history[sid][idx]

    return render_template(
        "test.html",
        error=None,
        generated_report=item["report_html"],
        file_ready=True,
        filename=item["filename"],
        history=search_history[sid],
    )

@app.route('/download')
@app.route('/download/<int:idx>')
def download_report(idx=0):
    """Download a specific report from history by index, defaulting to the latest (index 0)"""
    try:
        sid = get_session_id()
        if sid not in search_history or idx < 0 or idx >= len(search_history[sid]):
            # If the index is out of bounds, try to download the latest report
            if idx != 0 and len(search_history[sid]) > 0:
                idx = 0
            else:
                return "Report not found", 404
        
        # Get specific report data
        report_data = search_history[sid][idx]['report_dict']
        search_mode = search_history[sid][idx]['search_mode']

        # Create appropriate document based on search mode
        if search_mode == "Deep Search":
            research_data = DeepCompanyResearch(**report_data)
            doc_buffer = create_deep_word_document(research_data)
        else:
            research_data = QuickCompanyResearch(**report_data)
            doc_buffer = create_quick_word_document(research_data)

        # Create clean filename
        company_clean = research_data.company.replace(' ', '_').replace('/', '_').replace('\\', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        mode_suffix = search_mode.replace(' ', '_')
        filename = f"Company_Research_{company_clean}_{mode_suffix}_{timestamp}.docx"

        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        return f"Error generating document: {str(e)}", 500

@app.route("/api/research", methods=["POST"])
def api_research():
    """API endpoint for research"""
    try:
        data = request.get_json()
        company_name = data.get("company_name", "").strip()
        search_mode = data.get("search_mode", "Quick Search").strip()

        if not company_name:
            return jsonify({"error": "Company name is required"}), 400

        logger.info(f"📡 API: Starting {search_mode} for: {company_name}")
        
        # Execute appropriate search
        if search_mode == "Deep Search":
            research_data = deep_search(company_name)
        else:
            research_data = quick_search(company_name)
        
        result = research_data.model_dump()
        return jsonify(result)

    except Exception as e:
        logging.error(f"API research error: {str(e)}")
        return jsonify({"error": str(e)}), 500

import os

# Replace the entire main section with this Azure-compatible version for Dual Mode:
if __name__ == '__main__':
    # Check if running on Azure
    is_azure = os.environ.get('WEBSITE_SITE_NAME') is not None
    
    if not is_azure:
        # Only show detailed logs in local development
        print("\n" + "="*80)
        print("🚀 DUAL-MODE AI SALES RESEARCH TOOL - VERSION 2.2")
        print("="*80)
        
        # Check search engine API keys and provide status
        search_engines_available = []
        
        if TAVILY_API_KEY and TAVILY_API_KEY != "YOUR_TAVILY_API_KEY":
            search_engines_available.append("✅ Tavily (Primary)")
        else:
            print("⚠️  Tavily API key not configured - add to .env: TAVILY_API_KEY=your_key")
            
        if SERPAPI_API_KEY and SERPAPI_API_KEY != "YOUR_SERPAPI_API_KEY":
            search_engines_available.append("✅ SerpApi (Fallback 1)")
        else:
            print("⚠️  SerpApi API key not configured - add to .env: SERPAPI_API_KEY=your_key")
            
        if BRAVE_SEARCH_API_KEY and BRAVE_SEARCH_API_KEY != "YOUR_BRAVE_SEARCH_API_KEY":
            search_engines_available.append("✅ Brave Search (Fallback 2)")
        else:
            print("⚠️  Brave Search API key not configured - add to .env: BRAVE_SEARCH_API_KEY=your_key")
        
        if not search_engines_available:
            print("\n❌ CRITICAL: No search engines configured!")
            print("To use this application, configure at least one search API key:")
            print("1. Tavily: https://tavily.com (Recommended primary)")
            print("2. SerpApi: https://serpapi.com (Google search results)")
            print("3. Brave Search: https://brave.com/search/api (Alternative)")
            print("Add keys to your .env file and restart the application.")
            print("="*80)
            exit(1)
        else:
            print(f"\n🔧 Search Engines Configured: {len(search_engines_available)}")
            for engine in search_engines_available:
                print(f"   {engine}")
        
        print("\n🎯 Dual-Mode Features (Version 2.2):")
        print("   ⚡ Quick Search - 6-7 page reports (fast, essential insights)")
        print("   🔬 Deep Search - 15-20 page reports (exhaustive analysis + citations)")
        print("   ✨ Multi-tool search with intelligent fallback")
        print("   📊 Real-time data collection from multiple sources")
        print("   📅 Date freshness tracking")
        print("   🔄 Seamless fallback mechanisms")
        print("   📈 Enhanced data quality assurance")
        print("   🎨 Beautiful dual-mode reporting")
        print("   📋 Comprehensive metadata tracking")
        print("   📑 Source citations in Deep Search mode")
        
        print(f"\n🌐 Access the Enhanced Application at: http://localhost:5000")
        print("="*80)
    
    # Azure-compatible app startup
    if is_azure:
        # Azure production settings
        port = int(os.environ.get('PORT', 8000))
        app.run(debug=False, host='0.0.0.0', port=port)
    else:
        # Local development settings
        app.run(debug=True, host='127.0.0.1', port=5000)